from docx import Document
from PIL import Image, ImageDraw
import os

def create_color_image(filename, text, color):
    img = Image.new('RGB', (400, 300), color=color)
    d = ImageDraw.Draw(img)
    d.text((50, 140), text, fill=(255,255,255))
    img.save(filename)
    return filename

def generate_messy_inspection(filename):
    doc = Document()
    doc.add_heading('Field Notes - UrbanRoof Inspection', 0)
    
    doc.add_paragraph('Walked the site today (Nov 12). Met with the building mgr. Overall interior looks mostly fine but we found some weird stuff outside.')
    
    doc.add_paragraph('Notes on the North side:\n- Saw a huge crack running down the masonry wall near the loading bay door. Looks pretty bad, paint is peeling everywhere. \n- Tenant said it started leaking last week when it rained.')
    img_path = create_color_image('insp_crack.jpg', 'Camera View: Huge Masonry Crack', 'darkred')
    doc.add_picture(img_path, width=4000000)

    doc.add_paragraph('Notes on the Roof:\n- Actually could not access the East sector of the roof at all because the access door was jammed. So no data there.')
    doc.add_paragraph('- On the Southwest corner, the drain is totally blocked by leaves and mud. I see a big puddle. The membrane looks okay visually though, no tears that I could spot.')
    img_path = create_color_image('insp_puddle.jpg', 'Camera View: Muddy Puddle near Drain', 'blue')
    doc.add_picture(img_path, width=4000000)

    doc.add_paragraph('Server Room (Interior):\n- No water on the server room floor, it was totally dry. HVAC is making a very loud screeching sound however.')
    
    doc.save(filename)
    for f in ['insp_crack.jpg', 'insp_puddle.jpg']:
        if os.path.exists(f): os.remove(f)
    print(f"Generated {filename}")

def generate_messy_thermal(filename):
    doc = Document()
    doc.add_heading('IR Drone & Handheld Thermography Logs', 0)
    
    doc.add_paragraph('Drone swept the roof. Handheld IR used inside.')
    
    doc.add_paragraph('ROOF SWEEP:\n- The Southwest corner shows a massive localized thermal anomaly directly under where the visual drone saw a puddle. Measurements read -5°C delta. This is absolutely subsurface water pooling between the insulation layers, meaning the membrane is actually torn and flooded underneath even if it looks okay on top!!')
    img_path = create_color_image('therm_puddle.jpg', 'IR View: Deep pooling under membrane', 'darkblue')
    doc.add_picture(img_path, width=4000000)
    
    doc.add_paragraph('- Found some minor heat loss on the East sector roof vents, but overall it looks fine.')
    
    doc.add_paragraph('LOADING BAY WALL:\n- Handheld scan of the masonry crack on the North wall loading bay: No active moisture plumes detected inside the crack itself right now, but there is heavy thermal bridging happening. Insulation is missing behind the block over a 3ft radius.')
    
    doc.add_paragraph('SERVER ROOM AC:\n- The AC compressor casing is reading normal levels (around 30°C). Seems fine thermally.')
    img_path = create_color_image('therm_hvac.jpg', 'IR View: Server Room HVAC (Normal)', 'green')
    doc.add_picture(img_path, width=4000000)

    doc.save(filename)
    for f in ['therm_puddle.jpg', 'therm_hvac.jpg']:
        if os.path.exists(f): os.remove(f)
    print(f"Generated {filename}")

if __name__ == "__main__":
    generate_messy_inspection('Sample_Inspection_Raw.docx')
    generate_messy_thermal('Sample_Thermal_Raw.docx')
